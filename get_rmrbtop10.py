# !/usr/bin/env python
# -*- coding:utf-8 -*- 

import datetime
import lxml.html
from lxml import etree
import re
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt 

try:
    from urllib.request import urlopen, Request
except ImportError:
    from urllib2 import urlopen, Request

def today():
	today = datetime.datetime.today().date()
	return str(today)
	# return '2015-04-27'

def get_rmrb_top10(save_dir):
	url = r'http://paperpost.people.com.cn/rmrb-%s.html' %today()
	text = urlopen(url).read()

	pattern = re.compile(r'<a href=\"(.*?)\"')
	top10_urls = re.findall(pattern,text)

	for eachurl in top10_urls:
		try:
			html = lxml.html.parse(eachurl)
		except:
			continue

		res = html.xpath('//div[@id=\"ozoom\"]/p')
		title = html.xpath('//h1')[0].text
		author = html.xpath('//h4')[0].text
		newdoc = Document()


		# obj_styles = newdoc.styles
		# obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
		# obj_font = obj_charstyle.font
		# obj_font.size = Pt(16)
		# obj_font.name = u'方正仿宋_GBK'


		newdoc.add_heading(title,0)
		newdoc.add_paragraph(author)


		for node in res:
			# newdoc.add_paragraph(node.text,style = 'CommentsStyle')
			paragraph = newdoc.add_paragraph()
			paragraph_format = paragraph.paragraph_format
			paragraph_format.line_spacing = Pt(30)
			run = paragraph.add_run(node.text)
			font = run.font
			font.name = u'方正仿宋_GBK'
			font.size = Pt(16)

		newdoc.save(save_dir+title+r'.docx')


if __name__ == '__main__':
	save_dir =  r'/root/rmrb_top10/%s/'%(today())
	if not os.path.isdir(save_dir):
		os.mkdir(save_dir)

	get_rmrb_top10(save_dir)
